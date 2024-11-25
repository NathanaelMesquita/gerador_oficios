from docx import Document
import streamlit as st
from datetime import datetime

# Modelos de ofício
MODELOS = {
    "CPF": "modelo_cpf.docx",
    "IMEI": "modelo_imei.docx",
    "TERMINAL": "modelo_terminal.docx"
}

def preencher_modelo(template_path, context, output_path):
    """
    Preenche o modelo DOCX substituindo os placeholders pelos valores fornecidos.
    """
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            if f"{{{key}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{key}}}", value)

    # Verificar também placeholders em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    if f"{{{key}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{key}}}", value)

    doc.save(output_path)

def gerar_oficios(context, numeros_oficio, identificadores_selecionados, valores_identificadores):
    """
    Gera ofícios para as operadoras com base nos identificadores selecionados.
    """
    operadoras = ["Vivo", "Claro", "TIM"]
    arquivos_oficios = []

    for identificador, valores in zip(identificadores_selecionados, valores_identificadores):
        modelo_path = MODELOS[identificador]  # Seleciona o modelo correto
        context[identificador] = "\n".join([f"  - {v.strip()}" for v in valores if v.strip()])  # Preenche o placeholder do identificador

        for operadora, numero_oficio in zip(operadoras, numeros_oficio):
            # Atualizar contexto com operadora e número do ofício correspondente
            context["operadora"] = operadora
            context["numero_oficio"] = numero_oficio

            # Caminho para salvar o DOCX preenchido
            preenchido_path = f"oficio_{operadora}_{identificador}.docx"
            preencher_modelo(modelo_path, context, preenchido_path)
            arquivos_oficios.append(preenchido_path)

    return arquivos_oficios

# Interface Streamlit
st.title("Gerador de Ofícios")
st.header("Preencha os dados abaixo para gerar os ofícios:")

# Entradas do usuário
numero_oficio_vivo = st.text_input("Número do Ofício para Vivo", "001")
numero_oficio_claro = st.text_input("Número do Ofício para Claro", "002")
numero_oficio_tim = st.text_input("Número do Ofício para TIM", "003")

ano_atual = datetime.now().year
tipo_referencia = st.selectbox("Tipo de Referência", ["IP", "VPI"])
numero_referencia = st.text_input(f"Número do {tipo_referencia}", "12345")
ano_referencia = st.text_input("Ano do IP/VPI", str(ano_atual))
email = st.text_input("E-mail para Resposta", "delegado@exemplo.com")
nome_delegado = st.text_input("Nome do Delegado", "Delegado Exemplo")
nome_delegacia = st.text_input("Nome da Delegacia", "Delegacia Exemplo")

# Identificadores
st.subheader("Selecione os Identificadores:")
identificadores_selecionados = st.multiselect("Identificadores", ["CPF", "IMEI", "TERMINAL"])

valores_identificadores = []
for identificador in identificadores_selecionados:
    valores = st.text_area(f"Lista de {identificador}s (um por linha)", "")
    valores_identificadores.append(valores.splitlines())

# Data atual formatada
data_atual = datetime.now()
data_formatada = data_atual.strftime("%d de %B de %Y")

# Armazenar caminhos para exibição persistente dos botões de download
if "arquivos_gerados" not in st.session_state:
    st.session_state.arquivos_gerados = []

# Gerar Ofícios
if st.button("Gerar Ofícios"):
    # Criar o contexto para substituição
    context = {
        "ano_atual": str(ano_atual),
        "data_atual": data_atual.strftime("%d de %B de %Y"),
        "IP/VPI": tipo_referencia,
        "numero_ip_vpi": numero_referencia,
        "ano_ip_vpi": ano_referencia,
        "email_delegacia": email,
        "nome_delegado": nome_delegado,
        "nome_delegacia": nome_delegacia,
        # Inicializando os placeholders de identificadores
        "CPF": "",
        "IMEI": "",
        "TERMINAL": ""
    }

    # Gerar os documentos e salvar caminhos
    numeros_oficio = [numero_oficio_vivo, numero_oficio_claro, numero_oficio_tim]
    st.session_state.arquivos_gerados = gerar_oficios(context, numeros_oficio, identificadores_selecionados, valores_identificadores)

# Mostrar botões de download para os arquivos gerados
if st.session_state.arquivos_gerados:
    st.subheader("Baixar Ofícios Gerados:")
    for arquivo in st.session_state.arquivos_gerados:
        operadora_identificador = "_".join(arquivo.split('_')[1:3])  # Extrair operadora e identificador
        with open(arquivo, "rb") as docx_file:
            st.download_button(
                label=f"Baixar Ofício - {operadora_identificador}",
                data=docx_file,
                file_name=arquivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
